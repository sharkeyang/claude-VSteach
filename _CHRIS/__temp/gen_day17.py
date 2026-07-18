# -*- coding: utf-8 -*-
"""生成 Chris 7/17 每日计划 Word 版（一页紧凑版）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD  = RGBColor(0xB7, 0x8B, 0x4A)
GOLD2 = RGBColor(0xC9, 0xB9, 0x9A)
NAVY  = RGBColor(0x2D, 0x3E, 0x50)
BORDER = 'E8EAEB'
WHITE  = 'FFFFFF'

def set_font(run, font, size, color, bold=False):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rFonts.set(qn(a), font)

def cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)
    tcPr.append(shd)

def cell_borders(cell, color=BORDER, sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement('w:tcBorders')
    for edge in ('top','bottom'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz); b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcb.append(b)
    tcPr.append(tcb)

def cell_margins(cell, top=6, bottom=6, left=40, right=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge,val in (('top',top),('bottom',bottom),('left',left),('right',right)):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def row_height(row, h_pt):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'),str(int(h_pt*20))); h.set(qn('w:hRule'),'atLeast')
    trPr.append(h)

def fixed_layout(table):
    tblPr = table._tbl.tblPr
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)

def add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    return p

def run(p, text, font='微软雅黑', size=8, color=NAVY, bold=False):
    r = p.add_run(text); set_font(r, font, size, color, bold); return r

# ---------- 文档 ----------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.0)
sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)

# 1. 标题
p = add_para(doc, space_after=1)
run(p, 'C h r i s ', 'Georgia', 22, GOLD, True)
run(p, '每日计划', '微软雅黑', 12, NAVY, True)

# 2. 日期行
p = add_para(doc, space_after=2)
run(p, '──── 2026年7月17日 周五 · 第12天 ────', '微软雅黑', 9, NAVY)

# 3. 分隔线
p = add_para(doc, space_after=2)
run(p, '━'*49, '微软雅黑', 6, NAVY)

# ---------- 任务表 ----------
rows = [
    ('D12C1','🏀','6:50-8:00','1h10','篮球训练','运动'),
    ('','🍳','8:00-8:30','30m','早餐',''),
    ('D12C2','📜','8:30-9:00','30m','古文新篇《小石潭记》#26','古文'),
    ('','☕','9:00-9:10','10m','休息',''),
    ('D12C3','📘','9:10-9:55','45m','NC3 Lesson 14 新课+背诵','新概念'),
    ('','☕','9:55-10:05','10m','休息',''),
    ('D12C4','✍','10:05-10:35','30m','练字①《桃花源记》','练字'),
    ('','☕','10:35-10:45','10m','休息',''),
    ('D12C5','📚','10:45-12:30','1h45','学而思 9级 第16讲','数学'),
    ('','🍱','12:30-13:00','30m','午餐',''),
    ('D12C6','✍','13:00-13:30','30m','练字②《桃花源记》','练字'),
    ('','☕','13:30-13:40','10m','休息',''),
    ('D12C7','📜','13:40-14:00','20m','古文复习：生于忧患·次日','古文'),
    ('D12C8','🏃','17:00-19:00','2h','晚上活动','运动'),
    ('','🍽','19:00-19:30','30m','晚餐',''),
]

widths = [Cm(1.3), Cm(0.8), Cm(2.5), Cm(1.2), Cm(11.7)]
table = doc.add_table(rows=1, cols=5)
fixed_layout(table)
hdr = table.rows[0]
for i,w in enumerate(widths):
    hdr.cells[i].width = w

for (code,icon,time,dur,desc,cat) in rows:
    r = table.add_row()
    row_height(r, 13)
    cells = r.cells
    for i,w in enumerate(widths):
        cells[i].width = w
        cell_bg(cells[i], WHITE); cell_borders(cells[i]); cell_margins(cells[i])
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 编号
    p = cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, code, 'Georgia', 8, NAVY, True)
    # 图标
    p = cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, icon, 'Segoe UI Emoji', 9, GOLD)
    # 时间
    p = cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, time, 'Georgia', 8, NAVY)
    # 时长
    p = cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, dur, 'Georgia', 8, GOLD, True)
    # 内容
    p = cells[4].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(p, '□ ', '微软雅黑', 8, GOLD2)
    run(p, desc, '微软雅黑', 8, NAVY)
    if cat:
        run(p, '  ', '微软雅黑', 8, NAVY)
        run(p, cat, '微软雅黑', 7, GOLD)

# 第二条分隔线
p = add_para(doc, space_before=2, space_after=2)
run(p, '━'*49, '微软雅黑', 6, NAVY)

# ---------- 底部三栏 ----------
t2 = doc.add_table(rows=1, cols=3)
fixed_layout(t2)
w2 = [Cm(5.9), Cm(5.9), Cm(5.9)]
for i,w in enumerate(w2):
    c = t2.rows[0].cells[i]; c.width = w
    cell_bg(c, WHITE); cell_borders(c); cell_margins(c, top=6, bottom=6, left=60, right=60)
    c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def block(cell, title, lines):
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(p, title, '微软雅黑', 8, GOLD, True)
    for ln in lines:
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run(p, ln, '微软雅黑', 7.5, NAVY)

# 复习（记忆曲线）
block(t2.rows[0].cells[0], '复习  ｜  记忆曲线', [
    '📜 《生于忧患，死于安乐》· 次日',
    '📜 《记承天寺夜游》· 4天',
])
# 提问
block(t2.rows[0].cells[1], '提问', [
    'GD：小石潭记"潭中鱼可百许头"描绘了什么？',
    'GD：桃花源记核心思想？',
    'SX：学而思16讲新知识？',
])
# 完成
block(t2.rows[0].cells[2], '完成', [
    '□ 篮球  □ 小石潭记  □ NC3',
    '□ 练字  □ 学而思16讲  □ 复习',
    '□ 晚上活动',
])

# ---------- 备注区 ----------
p = add_para(doc, space_before=3, space_after=1)
run(p, '━'*49, '微软雅黑', 6, NAVY)

p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)
run(p, '✍ 练字要点', '微软雅黑', 8, GOLD, True)
for t in [
    '内容：《桃花源记》（7/16 背诵）→ 2 次×30min，≥2 页，质量>数量',
    '重点纠：中线对齐 · 大小一致 · 笔画力度（偏轻，需加重）',
    '写完拍照 → 发「Chris-练字-考试官」逐字评分',
]:
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
    p.paragraph_format.line_spacing = 1.05
    run(p, '· ' + t, '微软雅黑', 7.5, NAVY)

p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=1)
run(p, '📘 新概念要点', '微软雅黑', 8, GOLD, True)
for t in [
    'NC3 Lesson 14：新课学习 → 先听课文录音 → 逐句跟读 → 全文背诵',
    '昨天Lesson 13已背完，今天推进Lesson 14',
]:
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
    p.paragraph_format.line_spacing = 1.05
    run(p, '· ' + t, '微软雅黑', 7.5, NAVY)

out = r'D:\@VSwork\VSteach\_CHRIS\Chris每日计划20260717.docx'
doc.save(out)
print('SAVED:', out)
print('rows:', len(rows))