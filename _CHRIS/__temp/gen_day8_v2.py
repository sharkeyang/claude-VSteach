# -*- coding: utf-8 -*-
"""生成 Chris 7/14 每日计划 Word 版"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD  = RGBColor(0xB7, 0x8B, 0x4A)
NAVY  = RGBColor(0x2D, 0x3E, 0x50)
BORDER = 'E8EAEB'

def set_font(run, font, size, color, bold=False):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rFonts.set(qn(a), font)

def cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)
    tcPr.append(shd)

def cell_borders(cell, color=BORDER, sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement('w:tcBorders')
    for edge in ('top','bottom'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz); b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcb.append(b)
    tcPr.append(tcb)

def cell_margins(cell, top=6, bottom=6, left=40, right=40):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge,val in (('top',top),('bottom',bottom),('left',left),('right',right)):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def row_height(row, h_pt):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(h_pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def add_task_row(table, idx, icon, time, duration, content, category, n_cols=5):
    row = table.add_row()
    row_height(row, 13)
    cells = row.cells
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'D{idx}')
    set_font(r, 'Georgia', 8, NAVY, True)
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_borders(cells[0]); cell_margins(cells[0])
    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon)
    set_font(r, 'Segoe UI Emoji', 9, GOLD)
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_borders(cells[1]); cell_margins(cells[1])
    p = cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(time)
    set_font(r, 'Georgia', 8, NAVY)
    cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_borders(cells[2]); cell_margins(cells[2])
    p = cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(duration)
    set_font(r, 'Georgia', 8, GOLD, True)
    cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_borders(cells[3]); cell_margins(cells[3])
    p = cells[4].paragraphs[0]
    r = p.add_run(content)
    set_font(r, '微软雅黑', 8, NAVY)
    r2 = p.add_run(f'  {category}')
    set_font(r2, '微软雅黑', 7, GOLD)
    cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_borders(cells[4]); cell_margins(cells[4])

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('C h r i s')
set_font(r, 'Georgia', 22, GOLD, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('每日计划')
set_font(r, '微软雅黑', 12, NAVY, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('──── 2026年7月14日 周二 · 第8天 ────')
set_font(r, '微软雅黑', 9, NAVY)
for _ in range(3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━' * 49)
    set_font(r, '微软雅黑', 6, NAVY)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

col_widths = [Cm(1.3), Cm(0.8), Cm(2.5), Cm(1.2), Cm(11.7)]
table = doc.add_table(rows=0, cols=5)
table.autofit = False
for i, w in enumerate(col_widths):
    table.columns[i].width = w

tasks = [
    ('1', '🏀', '6:50-8:00', '1h10', '篮球训练', '运动'),
    ('', '🍳', '8:00-8:30', '30m', '早餐', ''),
    ('2', '📘', '8:30-9:15', '45m', 'NC3 Lesson 11 背诵', '新概念'),
    ('', '☕', '9:15-9:25', '10m', '休息', ''),
    ('3', '🗣', '9:25-10:55', '1h30', '发音训练官 · RE2 9A+B', '发音'),
    ('', '☕', '10:55-11:05', '10m', '休息', ''),
    ('4', '📜', '11:05-11:20', '15m', '古文 · 得道多助，失道寡助', '古文'),
    ('', '☕', '11:20-11:30', '10m', '休息', ''),
    ('5', '✍', '11:30-12:00', '30m', '练字①《记承天寺夜游》', '练字'),
    ('6', '🏀', '12:00-12:30', '30m', '篮球训练', '运动'),
    ('', '🍱', '12:30-13:00', '30m', '午餐', ''),
    ('7', '📚', '13:00-15:00', '2h', '学而思9级 第14讲 后半部分', '数学'),
    ('', '☕', '15:00-15:10', '10m', '休息', ''),
    ('8', '✍', '15:10-15:40', '30m', '练字②《记承天寺夜游》', '练字'),
    ('9', '🏃', '17:00-19:00', '2h', '运动', '运动'),
    ('', '🍽', '19:00-19:30', '30m', '晚餐', ''),
]

for idx, icon, time, dur, content, cat in tasks:
    add_task_row(table, idx, icon, time, dur, content, cat)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)

btable = doc.add_table(rows=1, cols=3)
for ci in range(3):
    cell = btable.rows[0].cells[ci]
    cell.width = Cm(5.9)
    cell_borders(cell, BORDER, '6')
    cell_margins(cell, 4, 4, 30, 30)

p = btable.cell(0,0).paragraphs[0]
r = p.add_run('【复习】')
set_font(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\n📜 《记承天寺夜游》· 次日\n📜 《得道多助》· 新')
set_font(r2, '微软雅黑', 7, NAVY)

p = btable.cell(0,1).paragraphs[0]
r = p.add_run('【提问】')
set_font(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\nGD："得道多助"中心论点？\nSX：第14讲核心思路？')
set_font(r2, '微软雅黑', 7, NAVY)

p = btable.cell(0,2).paragraphs[0]
r = p.add_run('【完成】')
set_font(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\n□ 篮球  □ NC3  □ 发音\n□ 古文  □ 练字  □ 学而思\n□ 运动')
set_font(r2, '微软雅黑', 7, NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run('✍ 练字要点：')
set_font(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('中线对齐 · 大小一致 · 笔画力度（偏轻需加重）· 拍照发「Chris-练字-考试官」评分')
set_font(r2, '微软雅黑', 7, NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
r = p.add_run('🗣 发音训练要点：')
set_font(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('用「Chris-英语发音-训练官」· 一句一句不打断 · 一次只纠1个问题 · RE2 9A+B')
set_font(r2, '微软雅黑', 7, NAVY)

out = r'D:\@VSwork\VSteach\_CHRIS\Chris每日计划20260714.docx'
doc.save(out)
print(f'[OK] 已保存: {out}')