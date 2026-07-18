# -*- coding: utf-8 -*-
"""生成 Chris 7/17 每日计划 Word 版"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD  = RGBColor(0xB7, 0x8B, 0x4A)
NAVY  = RGBColor(0x2D, 0x3E, 0x50)
BORDER = 'E8EAEB'

def sf(run, font, size, color, bold=False):
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rFonts.set(qn(a), font)

def cb(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement('w:tcBorders')
    for edge in ('top','bottom'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),BORDER)
        tcb.append(b)
    tcPr.append(tcb)

def cm(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge,val in (('top',6),('bottom',6),('left',40),('right',40)):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def rh(row, h_pt):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(h_pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def ar(table, idx, icon, time, dur, content, cat):
    row = table.add_row()
    rh(row, 13)
    cells = row.cells
    for ci in range(5):
        cb(cells[ci]); cm(cells[ci])
        cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if idx: r = p.add_run(f'D{idx}'); sf(r, 'Georgia', 8, NAVY, True)
    p = cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon); sf(r, 'Segoe UI Emoji', 9, GOLD)
    p = cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(time); sf(r, 'Georgia', 8, NAVY)
    p = cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(dur); sf(r, 'Georgia', 8, GOLD, True)
    p = cells[4].paragraphs[0]; r = p.add_run(content); sf(r, '微软雅黑', 8, NAVY)
    if cat: r2 = p.add_run(f'  {cat}'); sf(r2, '微软雅黑', 7, GOLD)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('C h r i s'); sf(r, 'Georgia', 22, GOLD, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('每日计划'); sf(r, '微软雅黑', 12, NAVY, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('──── 2026年7月17日 周五 ────'); sf(r, '微软雅黑', 9, NAVY)
for _ in range(3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━' * 49); sf(r, '微软雅黑', 6, NAVY)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

table = doc.add_table(rows=0, cols=5)
table.autofit = False
for i, w in enumerate([Cm(1.3), Cm(0.8), Cm(2.5), Cm(1.2), Cm(11.7)]):
    table.columns[i].width = w

tasks = [
    ('1', '🏀', '6:50-8:00', '1h10', '篮球训练', '运动'),
    ('', '🍳', '8:00-8:30', '30m', '早餐', ''),
    ('2', '📘', '8:30-9:15', '45m', 'NC3 Lesson 13 背诵', '新概念'),
    ('', '☕', '9:15-9:25', '10m', '休息', ''),
    ('3', '🗣', '9:25-10:55', '1h30', '发音训练官 · RE2 9A+B', '发音'),
    ('', '☕', '10:55-11:05', '10m', '休息', ''),
    ('4', '📜', '11:05-11:20', '15m', '古文 · 虽有嘉肴', '古文'),
    ('', '☕', '11:20-11:30', '10m', '休息', ''),
    ('5', '✍', '11:30-12:00', '30m', '练字①《北冥有鱼》', '练字'),
    ('6', '🏀', '12:00-12:30', '30m', '篮球训练', '运动'),
    ('', '🍱', '12:30-13:00', '30m', '午餐', ''),
    ('7', '📚', '13:00-15:00', '2h', '学而思9级 第15讲', '数学'),
    ('', '☕', '15:00-15:10', '10m', '休息', ''),
    ('8', '✍', '15:10-15:40', '30m', '练字②《北冥有鱼》', '练字'),
    ('9', '🏃', '17:00-19:00', '2h', '运动', '运动'),
    ('', '🍽', '19:00-19:30', '30m', '晚餐', ''),
]
for idx, icon, time, dur, content, cat in tasks:
    ar(table, idx, icon, time, dur, content, cat)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
bt = doc.add_table(rows=1, cols=3)
for ci in range(3):
    c = bt.rows[0].cells[ci]; c.width = Cm(5.9); cb(c); cm(c)

p = bt.cell(0,0).paragraphs[0]
r = p.add_run('【复习】'); sf(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\n📜 《北冥有鱼》· 次日\n📜 《虽有嘉肴》· 新'); sf(r2, '微软雅黑', 7, NAVY)

p = bt.cell(0,1).paragraphs[0]
r = p.add_run('【提问】'); sf(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\nGD："北冥有鱼"寓意？\nSX：第15讲内容？'); sf(r2, '微软雅黑', 7, NAVY)

p = bt.cell(0,2).paragraphs[0]
r = p.add_run('【完成】'); sf(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('\n□ 篮球  □ NC3  □ 发音\n□ 古文  □ 练字  □ 学而思\n□ 运动'); sf(r2, '微软雅黑', 7, NAVY)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
r = p.add_run('✍ 练字要点：'); sf(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('中线对齐 · 大小一致 · 笔画力度 · 拍照发「Chris-练字-考试官」评分'); sf(r2, '微软雅黑', 7, NAVY)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1)
r = p.add_run('🗣 发音训练要点：'); sf(r, '微软雅黑', 8, GOLD, True)
r2 = p.add_run('用「Chris-英语发音-训练官」· 一句一句不打断 · 一次只纠1个问题 · RE2 9A+B'); sf(r2, '微软雅黑', 7, NAVY)

out = r'D:\@VSwork\VSteach\_CHRIS\Chris每日计划20260717.docx'
doc.save(out)
print(f'[OK] 已保存: {out}')